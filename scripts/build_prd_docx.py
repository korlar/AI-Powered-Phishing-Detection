from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

SECTIONS = [
    (
        "Overview",
        [
            "The project will deliver a research-grade phishing detection system that classifies pasted email text and URLs as legitimate, spam, or phishing.",
            "The system uses Hugging Face Transformers with RoBERTa as the only fine-tuned model family. BERT is intentionally removed from scope to reduce training cost, memory usage, and serving complexity.",
            "The final deliverable is a client-server application featuring a robust FastAPI REST backend for inference and a Streamlit web frontend for user interaction.",
        ],
    ),
    (
        "Product Goals",
        [
            "Detect phishing intent in email body text using a fine-tuned RoBERTa classifier.",
            "Detect malicious or phishing URLs using a fine-tuned RoBERTa URL classifier.",
            "Provide a secure REST API (FastAPI) for programmatic access, batch processing, and token authentication.",
            "Provide a Dockerized Streamlit demo acting as a client to the backend API.",
        ],
    ),
    (
        "Non-Goals",
        [
            "Real-time Gmail, Outlook, or enterprise mail server integration.",
            "Multi-language detection beyond English.",
            "Production latency optimization, autoscaling, or cloud deployment hardening.",
            "Running BERT and RoBERTa simultaneously for comparison.",
        ],
    ),
]


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)


def add_table(document: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Light Shading Accent 1"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    document.add_paragraph()


def build_docx(output_path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Product Requirements Document")
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(31, 78, 121)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("AI-Powered Phishing Detection | RoBERTa-Only Implementation")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(80, 80, 80)

    meta = document.add_table(rows=3, cols=2)
    meta.style = "Light List Accent 1"
    for label, value in [
        ("Implementation stack", "Docker, VS Code, Python, Hugging Face Transformers, Streamlit"),
        ("Model decision", "Fine-tune RoBERTa only; no BERT comparator in implementation scope"),
        ("Document status", "Ready for implementation planning and codebase build-out"),
    ]:
        cells = meta.add_row().cells if label != "Implementation stack" else meta.rows[0].cells
        set_cell_text(cells[0], label, bold=True)
        set_cell_text(cells[1], value)

    for heading, bullets in SECTIONS:
        document.add_heading(heading, level=1)
        for bullet in bullets:
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(bullet)

    add_table(
        document,
        "Core User Stories",
        ["ID", "User Story", "Acceptance Criteria"],
        [
            [
                "US-01",
                "Paste an email body and get a verdict.",
                "Returns Legitimate, Spam, or Phishing with confidence.",
            ],
            [
                "US-02",
                "Paste a URL and get a verdict.",
                "Returns Legitimate or Phishing with confidence.",
            ],
            [
                "US-03",
                "Paste an email containing URLs.",
                "Intelligently routes the input for optimal classification.",
            ],
            ["US-04", "Fine-tune the email model from CSV data.", "Saves models/email-roberta."],
            ["US-05", "Fine-tune the URL model from CSV data.", "Saves models/url-roberta."],
            ["US-06", "Run the project in Docker.", "docker compose up --build starts Streamlit."],
            ["US-07", "Develop in VS Code.", ".vscode and .devcontainer settings are present."],
            [
                "US-08",
                "Authenticate to the API.",
                "The /token endpoint accepts credentials and returns a JWT.",
            ],
            [
                "US-09",
                "Submit a batch of URLs/texts.",
                "Processes bulk items and displays aggregate statistics.",
            ],
        ],
    )

    add_table(
        document,
        "Dataset Sources",
        ["Area", "Dataset Source", "Labels", "Purpose"],
        [
            ["Email", "CEAS-08 (CEAS_08.csv)", "Spam / Ham", "Spam and ham baseline training data"],
            [
                "Email",
                "Enron Corpus (Enron.csv)",
                "Ham",
                "Legitimate corporate email negative class",
            ],
            ["Email", "Ling-Spam (Ling.csv)", "Spam / Ham", "Spam/Legitimate baseline data"],
            [
                "Email",
                "Nazario Phishing (Nazario*.csv)",
                "Phishing",
                "Phishing email examples corpus",
            ],
            [
                "Email",
                "Nigerian Fraud (Nigerian*.csv)",
                "Phishing",
                "Advance-fee fraud email corpus",
            ],
            [
                "Email",
                "SpamAssassin (SpamAssasin.csv)",
                "Spam / Ham",
                "Public spam/legitimate corpus",
            ],
            ["Email", "TREC 2005/06/07 (TREC_0*.csv)", "Spam / Ham", "Spam track email datasets"],
            [
                "Email",
                "Hugging Face Hub (SetFit/enron_spam)",
                "Spam / Ham",
                "Standard split of Enron",
            ],
            [
                "Email",
                "Hugging Face Hub (Teddyha/phishing_benign)",
                "Benign / Phishing",
                "Phishing email examples",
            ],
            [
                "Email",
                "Hugging Face Hub (it4lia/PhishingEmailCurated)",
                "Benign / Phishing",
                "Cleaned phishing emails",
            ],
            [
                "Email",
                "Hugging Face Hub (luongnv89/phishing-email)",
                "Benign / Phishing",
                "Phishing email examples",
            ],
            [
                "Email",
                "Hugging Face Hub (Gunjand07/email-spam-dataset)",
                "Spam / Ham",
                "Modern spam corpus",
            ],
            [
                "Email",
                "Hugging Face Hub (SchoolP/Email_Spam_Dataset)",
                "Spam / Ham",
                "Supplemental spam emails",
            ],
            [
                "Email",
                "Hugging Face Hub (pleasenotagain/sanct-classify)",
                "Spam / Ham",
                "Supplemental spam emails",
            ],
            [
                "Email",
                "Hugging Face Hub (matefh/bitext-customer-support)",
                "Safe",
                "Non-phishing customer intents",
            ],
            [
                "Email",
                "Hugging Face Hub (SofienK-s/seven-phishing-email)",
                "Benign / Phishing",
                "Consolidated phishing dataset",
            ],
            [
                "Email",
                "Hugging Face Hub (JinqiangDing/seven-phishing-email)",
                "Benign / Phishing",
                "Consolidated phishing dataset",
            ],
            [
                "Email",
                "Hugging Face Hub (NatalieBob/phishing-email-dataset)",
                "Benign / Phishing",
                "Phishing email examples",
            ],
            [
                "Email",
                "Hugging Face Hub (deevyanshu/phishing_detection)",
                "Benign / Phishing",
                "Phishing email examples",
            ],
            [
                "Email",
                "Hugging Face Hub (prathxa/phishing_detection)",
                "Benign / Phishing",
                "Phishing email examples",
            ],
            [
                "Email",
                "Hugging Face Hub (shivahoody007/Phishing_Link_Pattern)",
                "Benign / Phishing",
                "Phishing links patterns",
            ],
            [
                "Email",
                "Majestic Million (majestic_million.csv)",
                "Benign",
                "Standard benign baseline domains",
            ],
            [
                "URL",
                "ISCX-URL-2016 (ISCX-URL-2016/)",
                "Benign / Phishing / Malware / Defacement / Spam",
                "22 CSV files containing multi-class URLs",
            ],
            [
                "URL",
                "Kaggle Malicious URLs (malicious_phish.csv)",
                "Benign / Phishing / Malware / Defacement",
                "Large 651k malicious URLs dataset",
            ],
            ["URL", "Tranco Top Domains (top-1m.csv)", "Benign", "Top 1M benign baseline domains"],
            [
                "URL",
                "Hugging Face Hub (flwrlabs/fed-phishing-urls)",
                "Benign / Phishing",
                "Collaborative URL phishing dataset",
            ],
            [
                "URL",
                "Hugging Face Hub (mahmoud0333/PhishingURLsANDBenign)",
                "Benign / Phishing",
                "Balanced URL dataset",
            ],
            [
                "URL",
                "Hugging Face Hub (pirocheto/phishing-url)",
                "Benign / Phishing",
                "Programmatic URL dataset",
            ],
            ["URL", "PhishTank (PhishTank.csv)", "Phishing", "Active verified phishing URLs feed"],
            [
                "URL",
                "URLhaus (urlhaus_recent.csv)",
                "Malware / Phishing",
                "Threat intelligence malware link feed",
            ],
        ],
    )

    add_table(
        document,
        "Functional Requirements",
        ["Requirement", "Description", "Priority"],
        [
            ["FR-01", "Parse raw email text and clean HTML/noisy formatting.", "Must"],
            ["FR-02", "Extract embedded URLs from email text.", "Must"],
            ["FR-03", "Normalize URL strings before training and inference.", "Must"],
            ["FR-04", "Fine-tune RoBERTa for 3-class email classification.", "Must"],
            ["FR-05", "Fine-tune RoBERTa for 2-class URL classification.", "Must"],
            [
                "FR-06",
                "Intelligently route requests to URL or Email models based on content.",
                "Must",
            ],
            ["FR-07", "Provide a Streamlit frontend client for manual testing.", "Must"],
            ["FR-08", "Provide Docker and VS Code setup files.", "Must"],
            ["FR-09", "Export evaluation metrics and confusion matrices.", "Should"],
            [
                "FR-10",
                "Add model explainability using Gradient Norm (Saliency) feature importances.",
                "Must",
            ],
            [
                "FR-11",
                "Implement a FastAPI backend with single and batch inference routes.",
                "Must",
            ],
            ["FR-12", "Implement JWT token-based authentication and payload size limits.", "Must"],
            [
                "FR-13",
                "Log single and batch predictions to an SQLite database history log.",
                "Must",
            ],
            ["FR-14", "Protect API endpoints using IP-based request rate limiting.", "Must"],
            [
                "FR-15",
                "Provide a dashboard tab in the Streamlit UI to view, refresh, and clear prediction logs.",
                "Must",
            ],
        ],
    )

    add_table(
        document,
        "Model Requirements",
        ["Model", "Classes", "Max Length", "Output"],
        [
            ["Email RoBERTa", "Legitimate, Spam, Phishing", "512 tokens", "models/email-roberta"],
            ["URL RoBERTa", "Legitimate, Phishing", "128 tokens", "models/url-roberta"],
        ],
    )

    add_table(
        document,
        "Evaluation Targets",
        ["Metric", "Target"],
        [
            ["Phishing precision", "> 0.92"],
            ["Phishing recall", "> 0.95"],
            ["Macro F1", "> 0.93"],
            ["False negative rate", "< 5%"],
            ["False positive rate", "< 3%"],
        ],
    )

    add_table(
        document,
        "Actual Evaluation Results - Email Model",
        ["Metric", "Value", "Legitimate Class", "Spam Class", "Phishing Class"],
        [
            ["Precision", "Macro Avg: 94%", "98%", "97%", "87%"],
            ["Recall", "Macro Avg: 94%", "99%", "85%", "97%"],
            ["F1-score", "Macro Avg: 94%", "99%", "91%", "92%"],
            ["Accuracy", "Overall: 94%", "-", "-", "-"],
        ],
    )

    add_table(
        document,
        "Actual Evaluation Results - URL Model",
        ["Metric", "Value", "Legitimate Class", "Phishing Class"],
        [
            ["Precision", "Macro Avg: 96%", "95%", "97%"],
            ["Recall", "Macro Avg: 96%", "97%", "95%"],
            ["F1-score", "Macro Avg: 96%", "96%", "96%"],
            ["Accuracy", "Overall: 96%", "-", "-"],
        ],
    )

    add_table(
        document,
        "System Architecture",
        ["Component", "Technology", "Responsibility"],
        [
            [
                "Frontend UI",
                "Streamlit",
                "Input form, batch upload, prediction history logs, and REST API client",
            ],
            [
                "Backend API",
                "FastAPI",
                "Token auth, rate limiting, routing, JSON logging, model lifecycle",
            ],
            [
                "Database",
                "SQLite (data/predictions.db)",
                "Persistent logging of prediction history, supporting query and delete operations",
            ],
            ["Email preprocessing", "Python, BeautifulSoup", "Clean email text and extract URLs"],
            [
                "URL preprocessing",
                "urllib, tldextract-ready design",
                "Normalize URLs and derive structural features",
            ],
            ["Email model", "RoBERTa + Transformers", "Email classification"],
            ["URL model", "RoBERTa + Transformers", "URL classification"],
            ["Container", "Docker / Docker Compose", "Reproducible setup for both API and UI"],
            ["IDE", "VS Code", "Development and debugging"],
        ],
    )

    document.add_heading("Inference Flow", level=1)
    inference_flow_steps = [
        "1. Authentication: User or programmatic client authenticates with the FastAPI backend (/token endpoint) to receive a JWT access token.",
        "2. Rate Limiting: Custom RateLimitingMiddleware checks incoming request IP addresses. Clients are limited to 100 requests per 60 seconds, with excess requests rejected with an HTTP 429 Too Many Requests status code.",
        "3. Submission: Client submits a single text or batch payload via Streamlit UI or direct API call, passing the JWT token.",
        "4. Payload Enforcer: FastAPI middleware enforces maximum payload size limits (e.g. 5MB) and logs incoming requests.",
        "5. Intelligent Routing & Allowlist Bypass: The system checks if the trimmed input is a pure URL. If so, it is routed to the URL model. If it contains text, embedded URLs are extracted and sent to the URL model, while the main text is sent to the Email model. A static allowlist (KNOWN_SAFE_DOMAINS) is checked to bypass model inference for trusted domains.",
        "6. Verdict Aggregation: If any extracted URL is classified as phishing with confidence >= 80%, the entire pipeline output is overridden as Phishing. Otherwise, the email classifier's output dictates the verdict. If only URLs are present, the highest-confidence URL classification is used.",
        "7. Explainability Mapping: If explainability is enabled, the system computes Gradient Norm (Saliency) feature importances for words in the email input.",
        "8. Database Logging: The system logs every prediction run to the local SQLite database (data/predictions.db), storing input content, prediction type, classification label, confidence level, boolean phishing flag, and reason text.",
        "9. Response rendering & History Querying: FastAPI returns a structured JSON response, and the Streamlit client renders results, bar charts, metrics, and highlight overlays. Users can also view, refresh, and clear past scans on the 'Prediction History' tab.",
    ]
    for step in inference_flow_steps:
        document.add_paragraph(step)
    document.add_paragraph()

    add_table(
        document,
        "Repository Structure",
        ["Path", "Purpose"],
        [
            [
                "backend/main.py",
                "FastAPI application entry point, lifespans, logging, and middlewares",
            ],
            ["backend/api/auth.py", "FastAPI router for authentication and token generation"],
            [
                "backend/api/predict.py",
                "FastAPI router for single and batch predictions, including history retrieval and clearing",
            ],
            [
                "backend/core/config.py",
                "Centralized system configurations (model paths, thresholds, safety allowlist, security keys)",
            ],
            [
                "backend/core/database.py",
                "SQLite helper functions to initialize/log/retrieve/clear prediction history",
            ],
            [
                "backend/core/inference.py",
                "Core model inference wrapper and gradient-based word explainability (saliency) handler",
            ],
            ["backend/core/rate_limiter.py", "IP-based request rate limiting middleware"],
            [
                "backend/core/security.py",
                "JWT token decoding, password hashing, and user authentication utilities",
            ],
            [
                "backend/schemas/predict.py",
                "Pydantic validation schemas for prediction requests, batch inputs, and response payloads",
            ],
            [
                "frontend/streamlit_app.py",
                "Streamlit web client implementing input forms, batch file uploads, explainability highlights, and log history dashboard",
            ],
            ["configs/training.yaml", "Hyperparameter configurations for email and URL models"],
            ["docs/PRD.md", "Product Requirements Document (Markdown)"],
            [
                "docs/PhishingDetection_PRD_RoBERTa.docx",
                "Compiled Product Requirements Document (Word format)",
            ],
            [
                "src/preprocess/clean_and_map.py",
                "Standalone script for cleaning, mapping labels, and balancing datasets",
            ],
            [
                "src/preprocess/dataset_builder.py",
                "Creates stratified train/validation/test splits",
            ],
            [
                "src/preprocess/download_hf_dataset.py",
                "Programmatic downloader for Hugging Face Hub datasets",
            ],
            [
                "src/preprocess/email_parser.py",
                "Text extractor parsing HTML to plain text and regex URL finder",
            ],
            [
                "src/preprocess/tokenize_datasets.py",
                "Tokenizes dataset splits and outputs PyTorch format datasets",
            ],
            [
                "src/preprocess/url_normalizer.py",
                "Normalizes URL strings and strips tracking parameters",
            ],
            [
                "src/preprocess/visualize_distribution.py",
                "Visualizes dataset distribution and split stats",
            ],
            [
                "src/models/train_roberta.py",
                "Main script to load tokenized datasets and fine-tune RoBERTa-base",
            ],
            [
                "src/models/evaluate_roberta.py",
                "Computes test-split evaluation metrics and confusion matrices",
            ],
            [
                "src/models/evaluation_report.py",
                "Generates summary classification report text file",
            ],
            [
                "src/models/generate_report.py",
                "Generates diagnostic plots and confusion matrix images",
            ],
            [
                "src/pipeline/aggregator.py",
                "Implementation of security aggregation and fallback rules",
            ],
            ["scripts/build_prd_docx.py", "Script to convert the PRD sections into Word format"],
            ["scripts/eda.py", "Script to compute class statistics, word counts, and top words"],
            ["scripts/test_api.py", "Quick API integration smoke-test script"],
            ["tests/", "Unit test suite covering backend API, aggregator, and preprocessing rules"],
            [
                "reports/",
                "Output directory containing generated confusion matrices and evaluation reports",
            ],
            ["data/", "Data folder containing raw, processed, and tokenized splits (gitignored)"],
            ["models/", "Saved local model weights checkpoints directory (gitignored)"],
        ],
    )

    add_table(
        document,
        "Implementation Milestones",
        ["Phase", "Deliverables", "Duration"],
        [
            ["1. Data collection", "Raw datasets downloaded and documented", "1 week"],
            ["2. Preprocessing", "Email parser, URL normalizer, train/val/test splits", "1 week"],
            ["3. Sanity baseline", "Dataset stats and simple reference metrics", "3 days"],
            ["4. Email RoBERTa", "Fine-tuned email checkpoint and metrics", "1 week"],
            ["5. URL RoBERTa", "Fine-tuned URL checkpoint and metrics", "1 week"],
            ["6. API Backend", "FastAPI, authentication, and batch processing", "4 days"],
            ["7. Combined pipeline", "Routing logic and end-to-end inference path", "4 days"],
            ["8. Streamlit Client", "UI with batch upload, auth, and metrics", "4 days"],
            ["9. Evaluation write-up", "Final metrics, error analysis, README", "4 days"],
        ],
    )

    add_table(
        document,
        "Risks And Mitigations",
        ["Risk", "Impact", "Mitigation"],
        [
            [
                "Few high-quality phishing emails",
                "High",
                "Use class weights, augmentation, and external held-out evaluation.",
            ],
            [
                "Dataset artifact memorization",
                "High",
                "Deduplicate, evaluate across sources, and inspect false negatives.",
            ],
            [
                "GPU memory limitations",
                "Medium",
                "Use smaller batch size, gradient accumulation, and max length caps.",
            ],
            [
                "URL strings lack context",
                "Medium",
                "Preserve structure and optionally add engineered URL features later.",
            ],
            ["App starts before models exist", "Low", "Show a clear missing-model warning."],
            [
                "API vulnerable to large payloads",
                "Medium",
                "Enforce maximum payload size limits via FastAPI middleware.",
            ],
            [
                "API abuse and Denial-of-Service",
                "High",
                "Implement client IP-based rate limiting (100 req/60s) via middleware.",
            ],
        ],
    )

    document.add_heading("Release Criteria", level=1)
    for item in [
        "Docker build completes.",
        "Unit tests pass.",
        "FastAPI backend and Streamlit app communicate successfully.",
        "Authentication successfully secures the API endpoints.",
        "Training scripts accept processed CSV inputs.",
        "RoBERTa-only design is reflected in code, README, and documentation.",
        "Final report includes separate email, URL, and combined pipeline metrics.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("Open Questions (Resolved)", level=1)
    for item in [
        "Model Checkpoints: Checkpoints are stored locally in the models/ directory. Direct integration with Hugging Face Hub can be enabled via environment tokens.",
        "Training Environment: The codebase is configured to run training either on local CUDA GPU or fall back to CPU.",
        "Live Lookups: The system features a static local allowlist (KNOWN_SAFE_DOMAINS) in the routing logic to bypass model inference for highly trusted domains. Active live lookup feeds are planned for future revisions.",
        "Explainability: Model explainability using Gradient Norm (Saliency) mapping is fully implemented, allowing users to see visual feedback of word importances directly in the Streamlit UI.",
    ]:
        document.add_paragraph(item, style="List Bullet")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


if __name__ == "__main__":
    build_docx(Path("docs/PhishingDetection_PRD_RoBERTa.docx"))
